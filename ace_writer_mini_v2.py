
import streamlit as st
import pandas as pd
import openai
from docx import Document
import io
import time

# ───────────────────────────────────────────────
# CONFIGURACIÓN INICIAL
# ───────────────────────────────────────────────
st.set_page_config(page_title="ACE Writer Mini V36", layout="wide")
st.title("✍️ ACE Writer Mini – Versión 36 Estable y Verificada")

# ───────────────────────────────────────────────
# ESTADO INICIAL
# ───────────────────────────────────────────────
for key in ["clave_ok", "referencias_completas", "referencias_incompletas", "subtema", "redaccion", "citadas"]:
    if key not in st.session_state:
        st.session_state[key] = [] if "referencias" in key else ""

# ───────────────────────────────────────────────
# PASO 0 – CLAVE OPENAI
# ───────────────────────────────────────────────
api_key = st.text_input("🔐 Clave OpenAI", type="password")
if api_key.startswith("sk-"):
    st.session_state["clave_ok"] = True
    st.success("✅ Clave válida")
else:
    st.warning("🔑 Ingresá una clave válida para continuar")

# ───────────────────────────────────────────────
# PASO 1 – PLANTILLA WORD
# ───────────────────────────────────────────────
st.subheader("Paso 1 – Cargar plantilla Word (.dotx)")
plantilla = st.file_uploader("📂 Plantilla Word", type=["dotx"])
if plantilla:
    doc = Document(plantilla)
    requeridos = ["Heading 1", "Heading 2", "Normal", "Reference"]
    encontrados = [s.name for s in doc.styles]
    validacion = [{"Estilo": s, "Presente": "✅" if s in encontrados else "❌"} for s in requeridos]
    st.dataframe(pd.DataFrame(validacion))

# ───────────────────────────────────────────────
# PASO 2 – REFERENCIAS CSV
# ───────────────────────────────────────────────
st.subheader("Paso 2 – Cargar tabla de referencias (.csv)")
archivo_csv = st.file_uploader("📄 Archivo .csv", type=["csv"])
selected_refs = []
completas, incompletas = [], []

if archivo_csv:
    df = pd.read_csv(archivo_csv)
    columnas_requeridas = ["Autores", "Año", "Título del artículo", "Journal"]
    if not all(col in df.columns for col in columnas_requeridas):
        st.error("❌ El archivo debe incluir columnas: Autores, Año, Título del artículo, Journal")
    else:
        for i, row in df.iterrows():
            try:
                ref = f"{row['Autores']} ({row['Año']}). {row['Título del artículo']}. {row['Journal']}."
                if "DOI" in df.columns and pd.notna(row["DOI"]):
                    ref += f" https://doi.org/{row['DOI']}"
                if all(pd.notna(row.get(c)) and str(row.get(c)).strip() != "" for c in columnas_requeridas):
                    completas.append(ref)
                else:
                    incompletas.append(ref)
            except Exception as e:
                incompletas.append(f"Error en fila {i}: {e}")

        selected_refs.extend(completas)
        st.success(f"✅ {len(completas)} referencias completas cargadas automáticamente.")
        if incompletas:
            st.warning(f"⚠️ {len(incompletas)} referencias incompletas detectadas.")
            seleccionar_todas = st.checkbox("Seleccionar todas las incompletas")
            for i, ref in enumerate(incompletas):
                if seleccionar_todas or st.checkbox(ref, key=f"incomp_{i}"):
                    selected_refs.append(ref)

# ───────────────────────────────────────────────
# PASO 3 – SUBTÍTULO DEL CAPÍTULO
# ───────────────────────────────────────────────
st.subheader("Paso 3 – Ingresá el subtítulo del subtema")
st.session_state["subtema"] = st.text_input("✏️ Subtema", value=st.session_state["subtema"])

# ───────────────────────────────────────────────
# PASO 4 – GENERAR REDACCIÓN CON GPT
# ───────────────────────────────────────────────
def redactar_con_gpt(subtema, capitulo, referencias, api_key):
    prompt = f"""Actuás como redactor científico del Proyecto eBooks ACE.
Tu tarea es redactar el subtema titulado **{subtema}**, que forma parte del capítulo **{capitulo}** del eBook ACE.

📌 Condiciones obligatorias:
– El texto debe tener mínimo **1500 palabras reales**
– Incluir al menos **1 recurso visual sugerido cada 500 palabras**
– Utilizar **solo** las referencias incluidas a continuación
– Cerrar con una sección de referencias en formato **APA 7**, solo con las fuentes citadas realmente en el cuerpo

📚 Lista de referencias válidas:
{chr(10).join(referencias)}

Redactá el texto directamente a continuación, en tono técnico claro, orientado a entrenadores profesionales, con ejemplos prácticos y subtítulos.
"""

    try:
        client = openai.OpenAI(api_key=api_key)
        with st.spinner("✍️ Redactando contenido inicial..."):
            r1 = client.chat.completions.create(
                model="gpt-4",
                messages=[
                    {"role": "system", "content": "Sos un redactor técnico de contenidos científicos sobre entrenamiento."},
                    {"role": "user", "content": prompt}
                ],
                temperature=0.7,
                max_tokens=3500
            )
        texto = r1.choices[0].message.content
        if len(texto.split()) >= 1500:
            return texto

        # Autoampliación si el texto es corto
        extend_prompt = f"Extendé este texto sin repetir ideas hasta superar las 1500 palabras:\n\n{texto}"
        with st.spinner("🔁 Solicitando ampliación automática..."):
            r2 = client.chat.completions.create(
                model="gpt-4",
                messages=[
                    {"role": "system", "content": "Sos un redactor técnico de contenidos científicos sobre entrenamiento."},
                    {"role": "user", "content": extend_prompt}
                ],
                temperature=0.7,
                max_tokens=3000
            )
        ampliacion = r2.choices[0].message.content
        return texto + "\n\n" + ampliacion
    except Exception as e:
        st.error(f"⚠️ Error al generar redacción: {str(e)}")
        return ""

# ───────────────────────────────────────────────
# PASO 5 – BOTÓN DE REDACCIÓN
# ───────────────────────────────────────────────
capitulo = "Capítulo personalizado desde interfaz"
if st.button("🚀 Generar redacción completa"):
    if st.session_state["clave_ok"] and st.session_state["subtema"] and selected_refs:
        texto = redactar_con_gpt(st.session_state["subtema"], capitulo, selected_refs, api_key)
        st.session_state["redaccion"] = texto

        # Detección de citas reales
        citas = []
        for ref in selected_refs:
            apellido = ref.split(",")[0]
            if apellido in texto:
                citas.append(ref)
        st.session_state["citadas"] = list(set(citas))

# ───────────────────────────────────────────────
# PASO 6 – MOSTRAR RESULTADO Y CONTADORES
# ───────────────────────────────────────────────
if st.session_state.get("redaccion"):
    st.subheader("🧾 Redacción generada")
    st.text_area("Texto", value=st.session_state["redaccion"], height=500)
    st.markdown(f"📊 Palabras: **{len(st.session_state['redaccion'].split())}**")
    st.markdown(f"📚 Citas utilizadas: **{len(st.session_state['citadas'])}**")

# ───────────────────────────────────────────────
# PASO 7 – EXPORTAR A WORD
# ───────────────────────────────────────────────
if st.session_state.get("redaccion"):
    if st.button("💾 Exportar a Word"):
        doc = Document(plantilla) if plantilla else Document()
        doc.add_heading(st.session_state["subtema"], level=1)
        doc.add_paragraph(st.session_state["redaccion"])
        doc.add_page_break()
        doc.add_heading("Referencias citadas", level=2)
        for ref in st.session_state["citadas"]:
            doc.add_paragraph(ref)
        ruta = "/mnt/data/ace_writer_mini_v36.docx"
        doc.save(ruta)
        with open(ruta, "rb") as f:
            st.download_button("📥 Descargar Word", data=f, file_name="ACEWriter_v36.docx")
